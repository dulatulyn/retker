from __future__ import annotations

from pathlib import Path

_SEV_RU = {1: "низкая", 2: "пониженная", 3: "средняя", 4: "высокая", 5: "критическая"}
_COMPLIANCE = ("Соответствие: Закон РК «О персональных данных и их защите» № 94-V. "
               "ИИН обезличены.")


def _docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


def _xlsx_available() -> bool:
    try:
        import xlsxwriter  # noqa: F401
        return True
    except Exception:
        return False


def to_docx(report: dict, path: str) -> str | None:
    if not _docx_available():
        return None
    from docx import Document
    from docx.shared import Pt

    s = report.get("sections", {})
    doc = Document()
    doc.add_heading(report.get("title", "Инцидент"), level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Инцидент {report.get('incident_id', '')}").bold = True

    doc.add_heading("Резюме", level=1)
    doc.add_paragraph(s.get("summary") or "")

    doc.add_heading("Хронология", level=1)
    for t in s.get("timeline", []):
        doc.add_paragraph(f"{t.get('ts', '')} · {t.get('label', '')}", style="List Bullet")

    if s.get("hypothesis"):
        doc.add_heading("Гипотеза атаки", level=1)
        doc.add_paragraph(s["hypothesis"])

    doc.add_heading("Рекомендации", level=1)
    for r in s.get("recommended_actions", []):
        doc.add_paragraph(r, style="List Bullet")

    if s.get("mitre"):
        doc.add_heading("MITRE ATT&CK", level=1)
        doc.add_paragraph(", ".join(s["mitre"]))

    alerts = s.get("alerts", [])
    if alerts:
        doc.add_heading("Алерты", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Детектор", "Категория", "Критичность", "Заголовок", "Сущность"]):
            hdr[i].text = h
        for a in alerts:
            row = table.add_row().cells
            row[0].text = str(a.get("detector", ""))
            row[1].text = str(a.get("category", ""))
            row[2].text = _SEV_RU.get(a.get("severity"), str(a.get("severity", "")))
            row[3].text = str(a.get("title", ""))
            row[4].text = str(a.get("entity", ""))

    foot = doc.add_paragraph()
    run = foot.add_run(_COMPLIANCE)
    run.italic = True
    run.font.size = Pt(9)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def to_xlsx(report: dict, path: str) -> str | None:
    if not _xlsx_available():
        return None
    import xlsxwriter

    s = report.get("sections", {})
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    wb = xlsxwriter.Workbook(path)
    bold = wb.add_format({"bold": True})
    wrap = wb.add_format({"text_wrap": True, "valign": "top"})
    head = wb.add_format({"bold": True, "bg_color": "#1e1e1e", "font_color": "#ffffff"})

    inc = wb.add_worksheet("Инцидент")
    inc.set_column(0, 0, 22)
    inc.set_column(1, 1, 80, wrap)
    rows = [
        ("Инцидент", report.get("incident_id", "")),
        ("Заголовок", report.get("title", "")),
        ("Резюме", s.get("summary", "")),
        ("Гипотеза", s.get("hypothesis", "")),
        ("MITRE", ", ".join(s.get("mitre", []))),
        ("Рекомендации", "; ".join(s.get("recommended_actions", []))),
    ]
    for i, (k, v) in enumerate(rows):
        inc.write(i, 0, k, bold)
        inc.write(i, 1, v)

    al = wb.add_worksheet("Алерты")
    cols = ["Детектор", "Категория", "Критичность", "Заголовок", "Сущность"]
    for c, h in enumerate(cols):
        al.write(0, c, h, head)
    al.set_column(0, 4, 22)
    for r, a in enumerate(s.get("alerts", []), start=1):
        al.write(r, 0, str(a.get("detector", "")))
        al.write(r, 1, str(a.get("category", "")))
        al.write(r, 2, _SEV_RU.get(a.get("severity"), str(a.get("severity", ""))))
        al.write(r, 3, str(a.get("title", "")))
        al.write(r, 4, str(a.get("entity", "")))

    tl = wb.add_worksheet("Хронология")
    tl.set_column(0, 0, 24)
    tl.set_column(1, 1, 80, wrap)
    tl.write(0, 0, "Время", head)
    tl.write(0, 1, "Событие", head)
    for r, t in enumerate(s.get("timeline", []), start=1):
        tl.write(r, 0, str(t.get("ts", "")))
        tl.write(r, 1, str(t.get("label", "")))

    wb.close()
    return path


def export(report: dict, out_dir: str, basename: str | None = None) -> dict:
    base = basename or report.get("incident_id") or "incident"
    d = Path(out_dir)
    docx_path = to_docx(report, str(d / f"{base}.docx"))
    xlsx_path = to_xlsx(report, str(d / f"{base}.xlsx"))
    return {"docx": docx_path, "xlsx": xlsx_path, "markdown": report.get("markdown")}


def _to_bytes(report: dict, suffix: str, writer) -> bytes | None:
    import os
    import tempfile
    fd, p = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    try:
        if writer(report, p) is None:
            return None
        return Path(p).read_bytes()
    finally:
        try:
            os.remove(p)
        except OSError:
            pass


def to_docx_bytes(report: dict) -> bytes | None:
    return _to_bytes(report, ".docx", to_docx)


def to_xlsx_bytes(report: dict) -> bytes | None:
    return _to_bytes(report, ".xlsx", to_xlsx)
